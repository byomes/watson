"""jobs/documents/word.py — Read and create Word .docx files."""
import logging
import re

log = logging.getLogger(__name__)


def read_docx(path: str) -> str:
    from docx import Document
    try:
        doc = Document(path)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception as exc:
        log.error("read_docx failed: %s", exc)
        return f"Error reading Word document: {exc}"


def create_docx(text: str, path: str) -> bool:
    from docx import Document
    try:
        doc = Document()
        for line in text.splitlines():
            doc.add_paragraph(line)
        doc.save(path)
        return True
    except Exception as exc:
        log.error("create_docx failed: %s", exc)
        return False


def run(message: str = None) -> str:
    if not message:
        return "Word document skill ready."
    match = re.search(r'[\w/~.-]+\.docx?', message, re.IGNORECASE)
    if not match:
        return "Word document skill ready. Provide a file path to read a document."
    path = match.group(0).replace("~", __import__("os").path.expanduser("~"))
    text = read_docx(path)
    if not text:
        return f"No text extracted from {path}."
    preview = text[:1000]
    suffix = f"\n\n[{len(text)} chars total]" if len(text) > 1000 else ""
    return f"Word document: {path}\n\n{preview}{suffix}"
